"""Export API — download scripts in various formats."""

import io
from pathlib import Path
from urllib.parse import quote

import yaml
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.database import get_db
from app.models import Project, Chapter, Adaptation

router = APIRouter()


def _make_content_disposition(filename: str) -> str:
    """Generate a properly encoded Content-Disposition attachment header value.

    Follows RFC 5987: uses filename* with UTF-8 percent-encoding for non-ASCII
    characters, plus an ASCII-only filename fallback for legacy clients.
    """
    # ASCII fallback: strip non-ASCII chars
    ascii_name = "".join(c if ord(c) < 128 and c.isprintable() else "_" for c in filename)
    # RFC 5987: filename*=UTF-8''percent-encoded-utf8-bytes
    encoded = quote(filename, safe="")
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{encoded}'


@router.get("/projects/{project_id}/export/markdown")
async def export_markdown(project_id: str, db: AsyncSession = Depends(get_db)):
    """Export the full script as a Markdown file."""
    result = await db.execute(
        select(Project)
        .where(Project.id == project_id)
        .options(
            selectinload(Project.chapters).selectinload(Chapter.adaptations),
        )
    )
    project = result.scalars().unique().first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    style = project.style

    # Build markdown content
    lines = []
    lines.append(f"# {project.title}")
    if project.author:
        lines.append(f"\n**原著作者**：{project.author}")
    lines.append(f"\n**改编风格**：{style}")
    lines.append(f"\n**生成时间**：{project.updated_at.isoformat() if project.updated_at else ''}")
    lines.append("\n---\n")

    for chapter in sorted(project.chapters, key=lambda c: c.chapter_num):
        lines.append(f"## 第{chapter.chapter_num}章 {chapter.title or ''}\n")
        # Read from adaptation for current style
        script = _get_script_for_style(chapter, style)
        if script:
            lines.append(script)
        else:
            lines.append("（待改编...）\n")
        lines.append("\n---\n")

    content = "\n".join(lines)
    safe_title = project.title.replace(" ", "_").replace("/", "_")[:50]

    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": _make_content_disposition(f"{safe_title}_剧本.md")},
    )


@router.get("/projects/{project_id}/export/txt")
async def export_txt(project_id: str, db: AsyncSession = Depends(get_db)):
    """Export the full script as a plain text file."""
    result = await db.execute(
        select(Project)
        .where(Project.id == project_id)
        .options(
            selectinload(Project.chapters).selectinload(Chapter.adaptations),
        )
    )
    project = result.scalars().unique().first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    style = project.style

    lines = []
    lines.append(f"{project.title}")
    if project.author:
        lines.append(f"原著：{project.author}")
    lines.append("=" * 50)
    lines.append("")

    for chapter in sorted(project.chapters, key=lambda c: c.chapter_num):
        lines.append(f"【第{chapter.chapter_num}章】{chapter.title or ''}")
        lines.append("-" * 40)
        script = _get_script_for_style(chapter, style)
        lines.append(script or "（待改编...）")
        lines.append("")
        lines.append("")

    content = "\n".join(lines)
    safe_title = project.title.replace(" ", "_").replace("/", "_")[:50]

    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": _make_content_disposition(f"{safe_title}_剧本.txt")},
    )


@router.get("/projects/{project_id}/export/docx")
async def export_docx(project_id: str, db: AsyncSession = Depends(get_db)):
    """Export the full script as a Word (.docx) file."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    result = await db.execute(
        select(Project)
        .where(Project.id == project_id)
        .options(
            selectinload(Project.chapters).selectinload(Chapter.adaptations),
        )
    )
    project = result.scalars().unique().first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    style = project.style

    doc = Document()

    # Title
    title_para = doc.add_heading(project.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    if project.author:
        meta = doc.add_paragraph(f"原著作者：{project.author}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    for chapter in sorted(project.chapters, key=lambda c: c.chapter_num):
        doc.add_heading(f"第{chapter.chapter_num}章  {chapter.title or ''}", level=1)

        script = _get_script_for_style(chapter, style)
        if script:
            for line in script.split("\n"):
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("（待改编...）")

        doc.add_page_break()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = project.title.replace(" ", "_").replace("/", "_")[:50]

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": _make_content_disposition(f"{safe_title}_剧本.docx")
        },
    )


@router.get("/projects/{project_id}/export/yaml")
async def export_yaml(project_id: str, db: AsyncSession = Depends(get_db)):
    """Export the full script as a structured YAML file."""
    result = await db.execute(
        select(Project)
        .where(Project.id == project_id)
        .options(
            selectinload(Project.chapters).selectinload(Chapter.adaptations),
        )
    )
    project = result.scalars().unique().first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    style = project.style

    # Build YAML structure
    chapters_data: list[dict] = []
    for chapter in sorted(project.chapters, key=lambda c: c.chapter_num):
        chapter_entry: dict = {
            "chapter_num": chapter.chapter_num,
            "title": chapter.title or "",
            "scenes": [],
        }

        # Try structured scenes first, fall back to prose text
        adapt = _get_adaptation_for_style(chapter, style)
        if adapt and adapt.scenes and adapt.scenes.get("structured_scenes"):
            chapter_entry["scenes"] = adapt.scenes["structured_scenes"]
        else:
            # Fallback: embed prose script as a single raw scene
            script = _get_script_for_style(chapter, style)
            if script:
                chapter_entry["scenes"] = [
                    {
                        "scene_num": 1,
                        "time": None,
                        "location": None,
                        "characters": [],
                        "stage_directions": [],
                        "dialogues": [],
                        "raw_prose": script,
                    }
                ]

        chapters_data.append(chapter_entry)

    yaml_data = {
        "project": {
            "title": project.title,
            "author": project.author or "",
            "style": project.style,
            "total_chapters": project.total_chapters,
            "generated_at": project.updated_at.isoformat() if project.updated_at else "",
        },
        "chapters": chapters_data,
    }

    yaml_content = yaml.dump(
        yaml_data,
        allow_unicode=True,
        sort_keys=False,
        default_flow_style=False,
        indent=2,
    )

    safe_title = project.title.replace(" ", "_").replace("/", "_")[:50]

    return StreamingResponse(
        io.BytesIO(yaml_content.encode("utf-8")),
        media_type="application/x-yaml; charset=utf-8",
        headers={
            "Content-Disposition": _make_content_disposition(f"{safe_title}_剧本.yaml")
        },
    )


def _get_adaptation_for_style(chapter, style: str):
    """Get the Adaptation ORM object for a specific style."""
    for a in (chapter.adaptations or []):
        if a.style == style:
            return a
    return None


def _get_script_for_style(chapter, style: str) -> str | None:
    """Get the adaptation script_text for a specific style, with backward compat fallback."""
    for a in (chapter.adaptations or []):
        if a.style == style and a.script_text:
            return a.script_text
    # Fall back to legacy chapter.script_text (for data before migration)
    return chapter.script_text
